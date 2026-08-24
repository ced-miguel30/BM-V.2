"""Desglose de coste de los últimos N desayunos registrados (no anulados)."""

from __future__ import annotations

import csv
import json
import os
from datetime import date
from pathlib import Path

from docx import Document
from docx.shared import Cm, Pt
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

N = 3
SRC = Path(os.environ["LOCALAPPDATA"]) / "BM-V2-local" / "data" / "datos_hotel.json"
ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "exports"
OUT.mkdir(parents=True, exist_ok=True)
TODAY = date.today().isoformat()


def _esc(text: object) -> str:
    s = "" if text is None else str(text)
    return s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def main() -> None:
    data = json.loads(SRC.read_text(encoding="utf-8"))
    prods = {p["id"]: p for p in data.get("productos") or []}
    recetas = {r["id"]: r for r in data.get("recetas") or []}

    desayunos = [
        d for d in (data.get("desayunos") or []) if not d.get("anulado")
    ]
    desayunos.sort(
        key=lambda x: (
            str(x.get("fecha") or ""),
            str(x.get("hora") or ""),
            str(x.get("id") or ""),
        )
    )
    ultimos = desayunos[-N:]
    if not ultimos:
        raise SystemExit("No hay desayunos no anulados.")

    resumen_rows: list[dict] = []
    detalle_rows: list[dict] = []
    receta_rows: list[dict] = []

    for d in ultimos:
        did = d["id"]
        fecha = d.get("fecha") or ""
        hora = (d.get("hora") or "")[:19]
        coste_total = float(d.get("coste_total") or 0)
        huespedes = d.get("num_huespedes")
        coste_pax = (
            round(coste_total / float(huespedes), 4)
            if huespedes and float(huespedes) > 0
            else ""
        )
        resumen_rows.append(
            {
                "desayuno_id": did,
                "fecha": fecha,
                "hora": hora,
                "registrado_por": d.get("registrado_por") or "",
                "num_huespedes": huespedes if huespedes is not None else "",
                "coste_total_eur": round(coste_total, 2),
                "coste_por_huesped_eur": coste_pax,
                "n_lineas_producto": len(d.get("lineas") or []),
                "n_lineas_detalle": len(d.get("lineas_detalle") or []),
                "n_recetas": len(d.get("registros_recetas") or []),
                "observaciones": (d.get("observaciones") or "")[:120],
            }
        )

        # Prefer lineas_detalle (trazabilidad receta); fallback lineas
        dets = d.get("lineas_detalle") or []
        if not dets:
            for ln in d.get("lineas") or []:
                pid = ln.get("producto_id")
                prod = prods.get(pid) or {}
                uni = prod.get("unidad")
                if isinstance(uni, dict):
                    uni = uni.get("value") or uni
                detalle_rows.append(
                    {
                        "desayuno_id": did,
                        "fecha": fecha,
                        "origen": "linea",
                        "producto_id": pid or "",
                        "producto": prod.get("nombre") or pid or "",
                        "unidad": uni or "",
                        "cantidad": ln.get("cantidad"),
                        "coste_eur": ln.get("coste"),
                        "es_extra": ln.get("es_extra"),
                        "receta_origen": "",
                        "categoria_receta": "",
                        "tipo_servicio": "desayuno",
                    }
                )
        else:
            for ln in dets:
                pid = ln.get("producto_id")
                prod = prods.get(pid) or {}
                uni = prod.get("unidad")
                if isinstance(uni, dict):
                    uni = uni.get("value") or uni
                rid = ln.get("receta_origen_id") or ""
                rnom = (recetas.get(rid) or {}).get("nombre") or rid
                detalle_rows.append(
                    {
                        "desayuno_id": did,
                        "fecha": fecha,
                        "origen": ln.get("origen") or "detalle",
                        "producto_id": pid or "",
                        "producto": prod.get("nombre") or pid or "",
                        "unidad": uni or "",
                        "cantidad": ln.get("cantidad"),
                        "coste_eur": round(float(ln.get("coste") or 0), 4),
                        "es_extra": "",
                        "receta_origen": rnom,
                        "categoria_receta": ln.get("categoria_receta_snapshot")
                        or ln.get("categoria_receta")
                        or "",
                        "tipo_servicio": ln.get("tipo_servicio") or "desayuno",
                    }
                )

        for rr in d.get("registros_recetas") or []:
            rid = rr.get("receta_id") or ""
            receta_rows.append(
                {
                    "desayuno_id": did,
                    "fecha": fecha,
                    "receta_id": rid,
                    "receta": rr.get("nombre_receta")
                    or (recetas.get(rid) or {}).get("nombre")
                    or rid,
                    "porciones": rr.get("porciones"),
                    "porciones_estandar_snapshot": rr.get("porciones_estandar_snapshot"),
                    "factor_aplicado": rr.get("factor_aplicado"),
                    "categoria": rr.get("categoria_receta_snapshot") or "",
                    "n_extras": len(rr.get("extras") or []),
                    "n_omisiones": len(rr.get("omisiones") or []),
                }
            )

    csv_resumen = OUT / f"desayunos_ultimos{N}_resumen_{TODAY}.csv"
    csv_det = OUT / f"desayunos_ultimos{N}_detalle_productos_{TODAY}.csv"
    csv_rec = OUT / f"desayunos_ultimos{N}_recetas_{TODAY}.csv"
    for path, rows in (
        (csv_resumen, resumen_rows),
        (csv_det, detalle_rows),
        (csv_rec, receta_rows),
    ):
        with path.open("w", newline="", encoding="utf-8-sig") as f:
            w = csv.DictWriter(f, fieldnames=list(rows[0].keys()))
            w.writeheader()
            w.writerows(rows)

    pdf_path = OUT / f"desayunos_ultimos{N}_desglose_{TODAY}.pdf"
    docx_path = OUT / f"desayunos_ultimos{N}_desglose_{TODAY}.docx"
    _pdf(resumen_rows, detalle_rows, receta_rows, pdf_path)
    _docx(resumen_rows, detalle_rows, receta_rows, docx_path)

    print("SOURCE", SRC)
    print("Desayunos:", [r["desayuno_id"] for r in resumen_rows])
    for r in resumen_rows:
        print(
            f"  {r['desayuno_id']} {r['fecha']} {r['hora']} "
            f"total={r['coste_total_eur']}€ pax={r['num_huespedes']}"
        )
    print("CSV resumen:", csv_resumen.resolve())
    print("CSV detalle:", csv_det.resolve())
    print("CSV recetas:", csv_rec.resolve())
    print("PDF:", pdf_path.resolve())
    print("DOCX:", docx_path.resolve())


def _pdf(resumen, detalle, recetas, path: Path) -> None:
    styles = getSampleStyleSheet()
    title = ParagraphStyle("T", parent=styles["Heading1"], fontSize=13, spaceAfter=6)
    h2 = ParagraphStyle("H", parent=styles["Heading2"], fontSize=10, spaceBefore=8, spaceAfter=4)
    body = ParagraphStyle("B", parent=styles["Normal"], fontSize=8, leading=10)
    cell = ParagraphStyle("C", parent=styles["Normal"], fontSize=6, leading=7.5)

    doc = SimpleDocTemplate(
        str(path),
        pagesize=landscape(A4),
        leftMargin=8 * mm,
        rightMargin=8 * mm,
        topMargin=8 * mm,
        bottomMargin=8 * mm,
    )
    story: list = []
    story.append(Paragraph(f"Desglose de coste — últimos {N} desayunos", title))
    story.append(
        Paragraph(
            f"Fecha exportación {TODAY}. Solo registros no anulados. "
            "Costes tomados del snapshot del registro (no se recalcula FIFO ahora).",
            body,
        )
    )

    story.append(Paragraph("1. Resumen", h2))
    hdr = [
        Paragraph(x, cell)
        for x in (
            "<b>ID</b>",
            "<b>Fecha</b>",
            "<b>Hora</b>",
            "<b>Por</b>",
            "<b>Huéspedes</b>",
            "<b>Total €</b>",
            "<b>€/pax</b>",
            "<b>#prod</b>",
            "<b>#recetas</b>",
        )
    ]
    data = [hdr]
    for r in resumen:
        data.append(
            [
                Paragraph(_esc(r["desayuno_id"]), cell),
                Paragraph(_esc(r["fecha"]), cell),
                Paragraph(_esc(r["hora"]), cell),
                Paragraph(_esc(r["registrado_por"]), cell),
                Paragraph(_esc(r["num_huespedes"]), cell),
                Paragraph(_esc(r["coste_total_eur"]), cell),
                Paragraph(_esc(r["coste_por_huesped_eur"]), cell),
                Paragraph(_esc(r["n_lineas_producto"]), cell),
                Paragraph(_esc(r["n_recetas"]), cell),
            ]
        )
    t = Table(
        data,
        colWidths=[18 * mm, 22 * mm, 32 * mm, 28 * mm, 20 * mm, 20 * mm, 18 * mm, 16 * mm, 18 * mm],
        repeatRows=1,
    )
    t.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1B3A4B")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#CCC")),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F5F7FA")]),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]
        )
    )
    story.append(t)

    story.append(Paragraph("2. Recetas servidas en cada desayuno", h2))
    hdr2 = [
        Paragraph(x, cell)
        for x in (
            "<b>Desayuno</b>",
            "<b>Fecha</b>",
            "<b>Receta</b>",
            "<b>Porciones</b>",
            "<b>Factor</b>",
            "<b>Cat.</b>",
            "<b>Extras</b>",
            "<b>Omisiones</b>",
        )
    ]
    data2 = [hdr2]
    for r in recetas:
        data2.append(
            [
                Paragraph(_esc(r["desayuno_id"]), cell),
                Paragraph(_esc(r["fecha"]), cell),
                Paragraph(_esc(r["receta"]), cell),
                Paragraph(_esc(r["porciones"]), cell),
                Paragraph(_esc(r["factor_aplicado"]), cell),
                Paragraph(_esc(r["categoria"]), cell),
                Paragraph(_esc(r["n_extras"]), cell),
                Paragraph(_esc(r["n_omisiones"]), cell),
            ]
        )
    t2 = Table(
        data2,
        colWidths=[20 * mm, 22 * mm, 70 * mm, 20 * mm, 18 * mm, 28 * mm, 16 * mm, 18 * mm],
        repeatRows=1,
    )
    t2.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#0F6E56")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#CCC")),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F3FAF7")]),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]
        )
    )
    story.append(t2)

    story.append(Paragraph("3. Detalle de productos / coste (lineas_detalle)", h2))
    hdr3 = [
        Paragraph(x, cell)
        for x in (
            "<b>Desayuno</b>",
            "<b>Producto</b>",
            "<b>Ud</b>",
            "<b>Cant.</b>",
            "<b>€</b>",
            "<b>Receta origen</b>",
            "<b>Origen</b>",
        )
    ]
    data3 = [hdr3]
    for r in detalle:
        data3.append(
            [
                Paragraph(_esc(r["desayuno_id"]), cell),
                Paragraph(_esc(r["producto"]), cell),
                Paragraph(_esc(r["unidad"]), cell),
                Paragraph(_esc(r["cantidad"]), cell),
                Paragraph(_esc(r["coste_eur"]), cell),
                Paragraph(_esc(r["receta_origen"]), cell),
                Paragraph(_esc(r["origen"]), cell),
            ]
        )
    t3 = Table(
        data3,
        colWidths=[18 * mm, 70 * mm, 14 * mm, 20 * mm, 18 * mm, 70 * mm, 22 * mm],
        repeatRows=1,
    )
    t3.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#7A5C00")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("GRID", (0, 0), (-1, -1), 0.2, colors.HexColor("#CCC")),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#FFF8E7")]),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]
        )
    )
    story.append(t3)

    # Totales por desayuno check
    story.append(Spacer(1, 6))
    story.append(Paragraph("Comprobación suma detalle vs coste_total", h2))
    for r in resumen:
        did = r["desayuno_id"]
        s = sum(float(x["coste_eur"] or 0) for x in detalle if x["desayuno_id"] == did)
        story.append(
            Paragraph(
                f"· {_esc(did)}: coste_total={_esc(r['coste_total_eur'])} € · "
                f"suma detalle={s:.2f} € · diff={float(r['coste_total_eur']) - s:.2f} €",
                body,
            )
        )

    doc.build(story)


def _docx(resumen, detalle, recetas, path: Path) -> None:
    document = Document()
    section = document.sections[0]
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.left_margin = Cm(1.0)
    section.right_margin = Cm(1.0)
    section.top_margin = Cm(1.0)
    section.bottom_margin = Cm(1.0)

    document.add_heading(f"Desglose de coste — últimos {N} desayunos", level=1)
    p = document.add_paragraph(
        f"Exportación {TODAY}. Registros no anulados. Costes del snapshot del desayuno."
    )
    for run in p.runs:
        run.font.size = Pt(9)

    document.add_heading("1. Resumen", level=2)
    t = document.add_table(rows=1, cols=9)
    t.style = "Table Grid"
    for i, h in enumerate(
        ["ID", "Fecha", "Hora", "Por", "Huéspedes", "Total €", "€/pax", "#prod", "#recetas"]
    ):
        t.rows[0].cells[i].text = h
    for r in resumen:
        row = t.add_row().cells
        row[0].text = str(r["desayuno_id"])
        row[1].text = str(r["fecha"])
        row[2].text = str(r["hora"])
        row[3].text = str(r["registrado_por"])
        row[4].text = str(r["num_huespedes"])
        row[5].text = str(r["coste_total_eur"])
        row[6].text = str(r["coste_por_huesped_eur"])
        row[7].text = str(r["n_lineas_producto"])
        row[8].text = str(r["n_recetas"])

    document.add_heading("2. Recetas", level=2)
    t2 = document.add_table(rows=1, cols=6)
    t2.style = "Table Grid"
    for i, h in enumerate(["Desayuno", "Fecha", "Receta", "Porciones", "Factor", "Cat."]):
        t2.rows[0].cells[i].text = h
    for r in recetas:
        row = t2.add_row().cells
        row[0].text = str(r["desayuno_id"])
        row[1].text = str(r["fecha"])
        row[2].text = str(r["receta"])
        row[3].text = str(r["porciones"])
        row[4].text = str(r["factor_aplicado"] or "")
        row[5].text = str(r["categoria"] or "")

    document.add_heading("3. Detalle productos / coste", level=2)
    t3 = document.add_table(rows=1, cols=7)
    t3.style = "Table Grid"
    for i, h in enumerate(
        ["Desayuno", "Producto", "Ud", "Cant.", "€", "Receta origen", "Origen"]
    ):
        t3.rows[0].cells[i].text = h
    for r in detalle:
        row = t3.add_row().cells
        row[0].text = str(r["desayuno_id"])
        row[1].text = str(r["producto"])
        row[2].text = str(r["unidad"])
        row[3].text = str(r["cantidad"])
        row[4].text = str(r["coste_eur"])
        row[5].text = str(r["receta_origen"])
        row[6].text = str(r["origen"])

    document.save(str(path))


if __name__ == "__main__":
    main()
