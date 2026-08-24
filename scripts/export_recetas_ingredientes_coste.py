"""Desglose de coste por receta e ingredientes individuales (porción estándar)."""

from __future__ import annotations

import csv
import os
import tempfile
from datetime import date, datetime
from pathlib import Path

os.environ["BM_TEST_ISOLATION"] = "1"

SRC = Path(os.environ["LOCALAPPDATA"]) / "BM-V2-local" / "data" / "datos_hotel.json"
TD = Path(tempfile.mkdtemp())
DST = TD / "datos_hotel.json"
DST.write_bytes(SRC.read_bytes())
os.environ["BM_DEMO_FILE"] = str(DST)

from app.bootstrap import configure_for_flet, get_container, reset_container
from app.core.auth.roles import ROL_DIRECCION
from app.core.auth.session import AuthSession, clear_test_session, set_test_session
from app.core.services import receta_service
from docx import Document
from docx.shared import Cm, Pt
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "exports"
OUT.mkdir(parents=True, exist_ok=True)
TODAY = date.today().isoformat()


def _esc(text: object) -> str:
    s = "" if text is None else str(text)
    return s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _fmt(n: object, d: int = 4) -> str:
    if n is None or n == "":
        return ""
    try:
        v = float(n)
    except (TypeError, ValueError):
        return str(n)
    s = f"{v:.{d}f}".rstrip("0").rstrip(".")
    return s


def main() -> None:
    reset_container()
    clear_test_session()
    configure_for_flet()
    set_test_session(
        AuthSession(
            True,
            "usuario",
            "u1",
            "Dir",
            ROL_DIRECCION,
            "s1",
            datetime.now().isoformat(timespec="seconds"),
        )
    )
    data = get_container().app_data_store.get()

    resumen_rows: list[dict] = []
    detalle_rows: list[dict] = []
    bloques: list[tuple[dict, list[dict]]] = []

    for rec in sorted(data.recetas or [], key=lambda x: (x.nombre or "").lower()):
        if not getattr(rec, "activo", True):
            continue
        cat = getattr(rec, "categoria", "")
        if hasattr(cat, "value"):
            cat = cat.value
        porc = float(getattr(rec, "porciones_estandar", 1) or 1)
        res = receta_service.valorar_receta(rec.id)

        resumen = {
            "receta_id": rec.id,
            "receta": rec.nombre,
            "categoria": cat or "",
            "porciones_estandar": porc,
            "n_ingredientes": len(getattr(rec, "ingredientes", None) or []),
            "coste_total_eur": res.coste_total if res.ok else "",
            "coste_por_racion_eur": res.coste_por_racion if res.ok else "",
            "coste_completo": res.coste_completo if res.ok else False,
            "ok": res.ok,
            "mensaje": "" if res.ok else (res.mensaje or "Sin valoración"),
        }
        resumen_rows.append(resumen)

        ings: list[dict] = []
        if res.ok:
            for ln in res.lineas or []:
                row = {
                    "receta_id": rec.id,
                    "receta": rec.nombre,
                    "categoria": cat or "",
                    "porciones_estandar": porc,
                    "producto_id": ln.producto_id,
                    "producto": ln.nombre,
                    "unidad": ln.unidad_nativa,
                    "cantidad_por_porcion": round(float(ln.cantidad_nativa or 0), 6),
                    "coste_unitario_eur": (
                        round(float(ln.coste_unitario_aplicable), 6)
                        if ln.coste_unitario_aplicable is not None
                        else ""
                    ),
                    "coste_ingrediente_eur": round(float(ln.coste_estimado or 0), 4),
                    "coste_incompleto": bool(ln.coste_incompleto),
                    "producto_inactivo": bool(ln.producto_inactivo),
                    "coste_total_receta_eur": res.coste_total,
                    "coste_por_racion_eur": res.coste_por_racion,
                }
                ings.append(row)
                detalle_rows.append(row)
        bloques.append((resumen, ings))

    csv_res = OUT / f"recetas_coste_resumen_{TODAY}.csv"
    csv_det = OUT / f"recetas_coste_ingredientes_{TODAY}.csv"
    with csv_res.open("w", newline="", encoding="utf-8-sig") as f:
        w = csv.DictWriter(f, fieldnames=list(resumen_rows[0].keys()))
        w.writeheader()
        w.writerows(resumen_rows)
    with csv_det.open("w", newline="", encoding="utf-8-sig") as f:
        w = csv.DictWriter(f, fieldnames=list(detalle_rows[0].keys()))
        w.writeheader()
        w.writerows(detalle_rows)

    pdf_path = OUT / f"recetas_coste_ingredientes_{TODAY}.pdf"
    docx_path = OUT / f"recetas_coste_ingredientes_{TODAY}.docx"
    _pdf(bloques, pdf_path)
    _docx(bloques, docx_path)

    ok_n = sum(1 for r in resumen_rows if r["ok"])
    print("SOURCE", SRC)
    print(f"Recetas activas: {len(resumen_rows)} (valoradas OK: {ok_n})")
    print("CSV resumen:", csv_res.resolve())
    print("CSV ingredientes:", csv_det.resolve())
    print("PDF:", pdf_path.resolve())
    print("DOCX:", docx_path.resolve())


def _pdf(bloques: list[tuple[dict, list[dict]]], path: Path) -> None:
    styles = getSampleStyleSheet()
    title = ParagraphStyle("T", parent=styles["Heading1"], fontSize=13, spaceAfter=6)
    h2 = ParagraphStyle("H", parent=styles["Heading2"], fontSize=9, spaceBefore=8, spaceAfter=3)
    body = ParagraphStyle("B", parent=styles["Normal"], fontSize=8, leading=10)
    cell = ParagraphStyle("C", parent=styles["Normal"], fontSize=6.5, leading=8)

    doc = SimpleDocTemplate(
        str(path),
        pagesize=landscape(A4),
        leftMargin=8 * mm,
        rightMargin=8 * mm,
        topMargin=8 * mm,
        bottomMargin=8 * mm,
    )
    story: list = [
        Paragraph("Desglose de coste — recetas e ingredientes", title),
        Paragraph(
            f"Fecha exportación {TODAY}. Coste teórico a porción estándar (FIFO actual). "
            "Cada bloque: receta → ingredientes con cantidad, €/ud y coste de esa línea.",
            body,
        ),
        Spacer(1, 4),
    ]

    for resumen, ings in bloques:
        head = (
            f"<b>{_esc(resumen['receta'])}</b> ({_esc(resumen['receta_id'])}) · "
            f"cat. {_esc(resumen['categoria'] or '—')} · "
            f"rendimiento {_fmt(resumen['porciones_estandar'], 2)} · "
        )
        if resumen["ok"]:
            head += (
                f"total <b>{_fmt(resumen['coste_total_eur'], 2)} €</b> · "
                f"€/ración <b>{_fmt(resumen['coste_por_racion_eur'], 4)} €</b>"
            )
            if not resumen["coste_completo"]:
                head += " · <font color='red'>coste incompleto</font>"
        else:
            head += f"<font color='red'>{_esc(resumen['mensaje'])}</font>"
        story.append(Paragraph(head, h2))

        if not ings:
            story.append(Paragraph("Sin ingredientes valorados.", body))
            continue

        data = [
            [
                Paragraph(x, cell)
                for x in (
                    "<b>Producto</b>",
                    "<b>Ud</b>",
                    "<b>Cant. / porción</b>",
                    "<b>€ / ud</b>",
                    "<b>€ línea</b>",
                    "<b>Nota</b>",
                )
            ]
        ]
        for ln in ings:
            nota = []
            if ln["coste_incompleto"]:
                nota.append("sin precio/lote")
            if ln["producto_inactivo"]:
                nota.append("inactivo")
            data.append(
                [
                    Paragraph(_esc(ln["producto"]), cell),
                    Paragraph(_esc(ln["unidad"]), cell),
                    Paragraph(_fmt(ln["cantidad_por_porcion"], 6), cell),
                    Paragraph(_fmt(ln["coste_unitario_eur"], 6), cell),
                    Paragraph(_fmt(ln["coste_ingrediente_eur"], 4), cell),
                    Paragraph(_esc(", ".join(nota)), cell),
                ]
            )
        t = Table(data, colWidths=[95 * mm, 18 * mm, 28 * mm, 28 * mm, 24 * mm, 40 * mm])
        t.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.Color(0.92, 0.92, 0.92)),
                    ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 2),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 2),
                    ("TOPPADDING", (0, 0), (-1, -1), 1),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 1),
                ]
            )
        )
        story.append(t)

    doc.build(story)


def _docx(bloques: list[tuple[dict, list[dict]]], path: Path) -> None:
    doc = Document()
    for s in doc.sections:
        s.page_width = Cm(29.7)
        s.page_height = Cm(21.0)
        s.left_margin = Cm(1.2)
        s.right_margin = Cm(1.2)
        s.top_margin = Cm(1.2)
        s.bottom_margin = Cm(1.2)

    t = doc.add_heading("Desglose de coste — recetas e ingredientes", level=1)
    t.runs[0].font.size = Pt(14)
    p = doc.add_paragraph(
        f"Fecha exportación {TODAY}. Coste teórico a porción estándar (FIFO actual)."
    )
    p.runs[0].font.size = Pt(9)

    for resumen, ings in bloques:
        h = doc.add_heading(str(resumen["receta"]), level=2)
        h.runs[0].font.size = Pt(11)
        meta = (
            f"{resumen['receta_id']} · {resumen['categoria'] or '—'} · "
            f"rendimiento {_fmt(resumen['porciones_estandar'], 2)}"
        )
        if resumen["ok"]:
            meta += (
                f" · total {_fmt(resumen['coste_total_eur'], 2)} €"
                f" · €/ración {_fmt(resumen['coste_por_racion_eur'], 4)} €"
            )
            if not resumen["coste_completo"]:
                meta += " · coste incompleto"
        else:
            meta += f" · {resumen['mensaje']}"
        mp = doc.add_paragraph(meta)
        mp.runs[0].font.size = Pt(8)

        if not ings:
            doc.add_paragraph("Sin ingredientes valorados.")
            continue

        table = doc.add_table(rows=1, cols=6)
        table.style = "Table Grid"
        hdrs = [
            "Producto",
            "Ud",
            "Cant. / porción",
            "€ / ud",
            "€ línea",
            "Nota",
        ]
        for i, htxt in enumerate(hdrs):
            table.rows[0].cells[i].text = htxt
        for ln in ings:
            nota = []
            if ln["coste_incompleto"]:
                nota.append("sin precio/lote")
            if ln["producto_inactivo"]:
                nota.append("inactivo")
            row = table.add_row().cells
            row[0].text = str(ln["producto"] or "")
            row[1].text = str(ln["unidad"] or "")
            row[2].text = _fmt(ln["cantidad_por_porcion"], 6)
            row[3].text = _fmt(ln["coste_unitario_eur"], 6)
            row[4].text = _fmt(ln["coste_ingrediente_eur"], 4)
            row[5].text = ", ".join(nota)
        doc.add_paragraph("")

    doc.save(str(path))


if __name__ == "__main__":
    main()
