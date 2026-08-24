"""Revisión de costes (solo nombres y euros): unitario, porción estándar, porciones registradas."""

from __future__ import annotations

import csv
import os
import tempfile
from datetime import date, datetime
from pathlib import Path

os.environ["BM_TEST_ISOLATION"] = "1"

SRC = Path(os.environ["LOCALAPPDATA"]) / "BM-V2-local" / "data" / "datos_hotel.json"
TD = Path(tempfile.mkdtemp())
DST = TD / "datos_hotel.json"
DST.write_bytes(SRC.read_bytes())
os.environ["BM_DEMO_FILE"] = str(DST)

from app.bootstrap import configure_for_flet, get_container, reset_container
from app.core.auth.roles import ROL_DIRECCION
from app.core.auth.session import AuthSession, clear_test_session, set_test_session
from app.core.services import receta_service
from app.core.services.inventory_batch_service import valorizar_cantidad_fifo
from docx import Document
from docx.shared import Cm, Pt
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "exports"
OUT.mkdir(parents=True, exist_ok=True)
TODAY = date.today().isoformat()
N_DESAYUNOS = 3


def _esc(text: object) -> str:
    s = "" if text is None else str(text)
    return s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _eur(n: object, d: int = 4) -> str:
    if n is None or n == "":
        return ""
    try:
        v = float(n)
    except (TypeError, ValueError):
        return ""
    return f"{v:.{d}f}".rstrip("0").rstrip(".")


def _write_csv(path: Path, rows: list[dict]) -> None:
    with path.open("w", newline="", encoding="utf-8-sig") as f:
        w = csv.DictWriter(f, fieldnames=list(rows[0].keys()))
        w.writeheader()
        w.writerows(rows)


def _table(data: list[list], widths: list[float]) -> Table:
    t = Table(data, colWidths=widths)
    t.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.Color(0.92, 0.92, 0.92)),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 3),
                ("RIGHTPADDING", (0, 0), (-1, -1), 3),
                ("TOPPADDING", (0, 0), (-1, -1), 2),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
            ]
        )
    )
    return t


def main() -> None:
    reset_container()
    clear_test_session()
    configure_for_flet()
    set_test_session(
        AuthSession(
            True,
            "usuario",
            "u1",
            "Dir",
            ROL_DIRECCION,
            "s1",
            datetime.now().isoformat(timespec="seconds"),
        )
    )
    data = get_container().app_data_store.get()
    raw = __import__("json").loads(SRC.read_text(encoding="utf-8"))

    # 1) Coste unitario producto
    unitario: list[dict] = []
    for p in sorted(data.productos, key=lambda x: (x.nombre or "").lower()):
        if not getattr(p, "activo", True):
            continue
        uni = p.unidad.value if hasattr(p.unidad, "value") else str(p.unidad)
        val = valorizar_cantidad_fifo(data, p.id, 1.0)
        cu = val.coste_unitario_aplicable
        if cu is None and val.coste and not val.incompleto:
            cu = float(val.coste)
        unitario.append(
            {
                "producto": p.nombre,
                "unidad": uni,
                "coste_unitario_eur": round(float(cu), 6) if cu is not None else "",
            }
        )

    # 2) Recetas: porción estándar + ingredientes
    recetas_resumen: list[dict] = []
    recetas_ings: list[dict] = []
    for rec in sorted(data.recetas or [], key=lambda x: (x.nombre or "").lower()):
        if not getattr(rec, "activo", True):
            continue
        res = receta_service.valorar_receta(rec.id)
        if not res.ok:
            continue
        recetas_resumen.append(
            {
                "receta": rec.nombre,
                "porciones_estandar": float(getattr(rec, "porciones_estandar", 1) or 1),
                "coste_total_eur": res.coste_total,
                "coste_por_porcion_estandar_eur": res.coste_por_racion,
            }
        )
        for ln in res.lineas or []:
            recetas_ings.append(
                {
                    "receta": rec.nombre,
                    "ingrediente": ln.nombre,
                    "cantidad": round(float(ln.cantidad_nativa or 0), 6),
                    "unidad": ln.unidad_nativa,
                    "coste_eur": round(float(ln.coste_estimado or 0), 4),
                }
            )

    # 3) Últimos 3 desayunos: porciones registradas (nombres + costes)
    desayunos = [d for d in (raw.get("desayunos") or []) if not d.get("anulado")]
    desayunos.sort(
        key=lambda x: (
            str(x.get("fecha") or ""),
            str(x.get("hora") or ""),
            str(x.get("id") or ""),
        )
    )
    ultimos = desayunos[-N_DESAYUNOS:]
    receta_nombre = {
        r.get("id"): r.get("nombre") for r in (raw.get("recetas") or [])
    }
    prod_nombre = {
        p.get("id"): p.get("nombre") for p in (raw.get("productos") or [])
    }

    des_resumen: list[dict] = []
    des_productos: list[dict] = []
    des_recetas: list[dict] = []

    for d in ultimos:
        fecha = d.get("fecha") or ""
        hora = (d.get("hora") or "")[:19]
        coste = float(d.get("coste_total") or 0)
        pax = d.get("num_huespedes")
        des_resumen.append(
            {
                "fecha": fecha,
                "hora": hora,
                "huespedes": pax if pax is not None else "",
                "coste_total_eur": round(coste, 2),
                "coste_por_huesped_eur": (
                    round(coste / float(pax), 4)
                    if pax and float(pax) > 0
                    else ""
                ),
            }
        )

        # Agregar porciones por receta (suma de porciones registradas)
        por_rec: dict[str, float] = {}
        for rr in d.get("registros_recetas") or []:
            nom = rr.get("nombre_receta") or receta_nombre.get(rr.get("receta_id")) or ""
            try:
                por = float(rr.get("porciones") or 0)
            except (TypeError, ValueError):
                por = 0.0
            if nom:
                por_rec[nom] = por_rec.get(nom, 0.0) + por

        # Coste por receta desde lineas_detalle
        coste_rec: dict[str, float] = {}
        dets = d.get("lineas_detalle") or []
        if dets:
            for ln in dets:
                rid = ln.get("receta_origen_id") or ""
                nom = receta_nombre.get(rid) or "(sin receta)"
                coste_rec[nom] = coste_rec.get(nom, 0.0) + float(ln.get("coste") or 0)
                pname = prod_nombre.get(ln.get("producto_id")) or ln.get("producto_id") or ""
                des_productos.append(
                    {
                        "fecha": fecha,
                        "producto": pname,
                        "cantidad": ln.get("cantidad"),
                        "coste_eur": round(float(ln.get("coste") or 0), 4),
                        "receta": nom if nom != "(sin receta)" else "",
                    }
                )
        else:
            for ln in d.get("lineas") or []:
                pname = prod_nombre.get(ln.get("producto_id")) or ""
                des_productos.append(
                    {
                        "fecha": fecha,
                        "producto": pname,
                        "cantidad": ln.get("cantidad"),
                        "coste_eur": round(float(ln.get("coste") or 0), 4),
                        "receta": "",
                    }
                )

        nombres = sorted(set(por_rec) | set(coste_rec))
        for nom in sorted(nombres, key=lambda x: x.lower()):
            des_recetas.append(
                {
                    "fecha": fecha,
                    "receta": nom,
                    "porciones_registradas": round(por_rec.get(nom, 0.0), 4) or "",
                    "coste_eur": round(coste_rec.get(nom, 0.0), 4),
                }
            )

    # CSV
    csv_u = OUT / f"revision_coste_unitario_{TODAY}.csv"
    csv_rs = OUT / f"revision_recetas_porcion_estandar_{TODAY}.csv"
    csv_ri = OUT / f"revision_recetas_ingredientes_{TODAY}.csv"
    csv_dr = OUT / f"revision_desayunos_ultimos3_resumen_{TODAY}.csv"
    csv_dp = OUT / f"revision_desayunos_ultimos3_productos_{TODAY}.csv"
    csv_drec = OUT / f"revision_desayunos_ultimos3_recetas_{TODAY}.csv"
    _write_csv(csv_u, unitario)
    _write_csv(csv_rs, recetas_resumen)
    _write_csv(csv_ri, recetas_ings)
    _write_csv(csv_dr, des_resumen)
    _write_csv(csv_dp, des_productos)
    _write_csv(csv_drec, des_recetas)

    pdf_path = OUT / f"revision_costes_nombres_{TODAY}.pdf"
    docx_path = OUT / f"revision_costes_nombres_{TODAY}.docx"
    _pdf(unitario, recetas_resumen, recetas_ings, des_resumen, des_productos, des_recetas, pdf_path)
    _docx(unitario, recetas_resumen, recetas_ings, des_resumen, des_productos, des_recetas, docx_path)

    print("SOURCE", SRC)
    print("PDF", pdf_path.resolve())
    print("DOCX", docx_path.resolve())
    for p in (csv_u, csv_rs, csv_ri, csv_dr, csv_dp, csv_drec):
        print("CSV", p.resolve())


def _pdf(unitario, recetas_resumen, recetas_ings, des_resumen, des_productos, des_recetas, path: Path) -> None:
    styles = getSampleStyleSheet()
    title = ParagraphStyle("T", parent=styles["Heading1"], fontSize=13, spaceAfter=6)
    h2 = ParagraphStyle("H", parent=styles["Heading2"], fontSize=10, spaceBefore=10, spaceAfter=4)
    body = ParagraphStyle("B", parent=styles["Normal"], fontSize=8, leading=10)
    cell = ParagraphStyle("C", parent=styles["Normal"], fontSize=7, leading=8.5)

    doc = SimpleDocTemplate(
        str(path),
        pagesize=landscape(A4),
        leftMargin=10 * mm,
        rightMargin=10 * mm,
        topMargin=10 * mm,
        bottomMargin=10 * mm,
    )
    story: list = [
        Paragraph("Revisión de costes (solo nombres y euros)", title),
        Paragraph(
            f"Exportación {TODAY}. Sin IDs ni categorías. "
            "1) Coste unitario · 2) Porción estándar · 3) Últimos 3 desayunos registrados.",
            body,
        ),
    ]

    story.append(Paragraph("1. Coste unitario de producto", h2))
    data = [[Paragraph(x, cell) for x in ("<b>Producto</b>", "<b>Unidad</b>", "<b>€ / ud</b>")]]
    for r in unitario:
        data.append(
            [
                Paragraph(_esc(r["producto"]), cell),
                Paragraph(_esc(r["unidad"]), cell),
                Paragraph(_eur(r["coste_unitario_eur"], 6), cell),
            ]
        )
    story.append(_table(data, [140 * mm, 30 * mm, 30 * mm]))

    story.append(Paragraph("2. Recetas — coste por porción estándar", h2))
    data = [
        [
            Paragraph(x, cell)
            for x in (
                "<b>Receta</b>",
                "<b>Porciones estándar</b>",
                "<b>Coste total €</b>",
                "<b>€ / porción</b>",
            )
        ]
    ]
    for r in recetas_resumen:
        data.append(
            [
                Paragraph(_esc(r["receta"]), cell),
                Paragraph(_eur(r["porciones_estandar"], 2), cell),
                Paragraph(_eur(r["coste_total_eur"], 2), cell),
                Paragraph(_eur(r["coste_por_porcion_estandar_eur"], 4), cell),
            ]
        )
    story.append(_table(data, [110 * mm, 35 * mm, 30 * mm, 30 * mm]))

    story.append(Paragraph("2b. Ingredientes por receta (porción estándar)", h2))
    data = [
        [
            Paragraph(x, cell)
            for x in (
                "<b>Receta</b>",
                "<b>Ingrediente</b>",
                "<b>Cantidad</b>",
                "<b>Ud</b>",
                "<b>€</b>",
            )
        ]
    ]
    for r in recetas_ings:
        data.append(
            [
                Paragraph(_esc(r["receta"]), cell),
                Paragraph(_esc(r["ingrediente"]), cell),
                Paragraph(_eur(r["cantidad"], 6), cell),
                Paragraph(_esc(r["unidad"]), cell),
                Paragraph(_eur(r["coste_eur"], 4), cell),
            ]
        )
    story.append(_table(data, [70 * mm, 70 * mm, 25 * mm, 20 * mm, 25 * mm]))

    story.append(Paragraph("3. Últimos 3 desayunos — resumen", h2))
    data = [
        [
            Paragraph(x, cell)
            for x in (
                "<b>Fecha</b>",
                "<b>Hora</b>",
                "<b>Huéspedes</b>",
                "<b>Coste total €</b>",
                "<b>€ / huésped</b>",
            )
        ]
    ]
    for r in des_resumen:
        data.append(
            [
                Paragraph(_esc(r["fecha"]), cell),
                Paragraph(_esc(r["hora"]), cell),
                Paragraph(_esc(r["huespedes"]), cell),
                Paragraph(_eur(r["coste_total_eur"], 2), cell),
                Paragraph(_eur(r["coste_por_huesped_eur"], 4), cell),
            ]
        )
    story.append(_table(data, [35 * mm, 50 * mm, 30 * mm, 35 * mm, 35 * mm]))

    story.append(Paragraph("3b. Porciones registradas (por receta)", h2))
    data = [
        [
            Paragraph(x, cell)
            for x in (
                "<b>Fecha</b>",
                "<b>Receta</b>",
                "<b>Porciones registradas</b>",
                "<b>Coste €</b>",
            )
        ]
    ]
    for r in des_recetas:
        data.append(
            [
                Paragraph(_esc(r["fecha"]), cell),
                Paragraph(_esc(r["receta"]), cell),
                Paragraph(_eur(r["porciones_registradas"], 4), cell),
                Paragraph(_eur(r["coste_eur"], 4), cell),
            ]
        )
    story.append(_table(data, [35 * mm, 100 * mm, 40 * mm, 30 * mm]))

    story.append(Paragraph("3c. Productos consumidos en esos desayunos", h2))
    data = [
        [
            Paragraph(x, cell)
            for x in (
                "<b>Fecha</b>",
                "<b>Producto</b>",
                "<b>Cantidad</b>",
                "<b>Coste €</b>",
                "<b>Receta</b>",
            )
        ]
    ]
    for r in des_productos:
        data.append(
            [
                Paragraph(_esc(r["fecha"]), cell),
                Paragraph(_esc(r["producto"]), cell),
                Paragraph(_eur(r["cantidad"], 6), cell),
                Paragraph(_eur(r["coste_eur"], 4), cell),
                Paragraph(_esc(r["receta"]), cell),
            ]
        )
    story.append(_table(data, [30 * mm, 70 * mm, 25 * mm, 25 * mm, 55 * mm]))

    doc.build(story)


def _docx(unitario, recetas_resumen, recetas_ings, des_resumen, des_productos, des_recetas, path: Path) -> None:
    doc = Document()
    for s in doc.sections:
        s.page_width = Cm(29.7)
        s.page_height = Cm(21.0)
        s.left_margin = Cm(1.2)
        s.right_margin = Cm(1.2)
        s.top_margin = Cm(1.2)
        s.bottom_margin = Cm(1.2)

    doc.add_heading("Revisión de costes (solo nombres y euros)", level=1)
    doc.add_paragraph(
        f"Exportación {TODAY}. Sin IDs ni categorías. "
        "Coste unitario · porción estándar · últimos 3 desayunos."
    ).runs[0].font.size = Pt(9)

    def add_table(headers: list[str], rows: list[list[str]]) -> None:
        t = doc.add_table(rows=1, cols=len(headers))
        t.style = "Table Grid"
        for i, h in enumerate(headers):
            t.rows[0].cells[i].text = h
        for row in rows:
            cells = t.add_row().cells
            for i, val in enumerate(row):
                cells[i].text = val
        doc.add_paragraph("")

    doc.add_heading("1. Coste unitario de producto", level=2)
    add_table(
        ["Producto", "Unidad", "€ / ud"],
        [[r["producto"], str(r["unidad"]), _eur(r["coste_unitario_eur"], 6)] for r in unitario],
    )

    doc.add_heading("2. Recetas — coste por porción estándar", level=2)
    add_table(
        ["Receta", "Porciones estándar", "Coste total €", "€ / porción"],
        [
            [
                r["receta"],
                _eur(r["porciones_estandar"], 2),
                _eur(r["coste_total_eur"], 2),
                _eur(r["coste_por_porcion_estandar_eur"], 4),
            ]
            for r in recetas_resumen
        ],
    )

    doc.add_heading("2b. Ingredientes por receta", level=2)
    add_table(
        ["Receta", "Ingrediente", "Cantidad", "Ud", "€"],
        [
            [
                r["receta"],
                r["ingrediente"],
                _eur(r["cantidad"], 6),
                str(r["unidad"] or ""),
                _eur(r["coste_eur"], 4),
            ]
            for r in recetas_ings
        ],
    )

    doc.add_heading("3. Últimos 3 desayunos — resumen", level=2)
    add_table(
        ["Fecha", "Hora", "Huéspedes", "Coste total €", "€ / huésped"],
        [
            [
                r["fecha"],
                r["hora"],
                str(r["huespedes"]),
                _eur(r["coste_total_eur"], 2),
                _eur(r["coste_por_huesped_eur"], 4),
            ]
            for r in des_resumen
        ],
    )

    doc.add_heading("3b. Porciones registradas (por receta)", level=2)
    add_table(
        ["Fecha", "Receta", "Porciones registradas", "Coste €"],
        [
            [
                r["fecha"],
                r["receta"],
                _eur(r["porciones_registradas"], 4),
                _eur(r["coste_eur"], 4),
            ]
            for r in des_recetas
        ],
    )

    doc.add_heading("3c. Productos consumidos", level=2)
    add_table(
        ["Fecha", "Producto", "Cantidad", "Coste €", "Receta"],
        [
            [
                r["fecha"],
                r["producto"],
                _eur(r["cantidad"], 6),
                _eur(r["coste_eur"], 4),
                r["receta"],
            ]
            for r in des_productos
        ],
    )

    doc.save(str(path))


if __name__ == "__main__":
    main()
