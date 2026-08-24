"""Desglose imprimible: coste unitario producto vs coste por porción estándar.

Ejemplo: lata champiñón 3 kg = €/Ud del envase; en desayuno se usa una fracción → €/ración.
"""

from __future__ import annotations

import csv
import os
import tempfile
from datetime import date, datetime
from pathlib import Path

os.environ["BM_TEST_ISOLATION"] = "1"

SRC = Path(os.environ["LOCALAPPDATA"]) / "BM-V2-local" / "data" / "datos_hotel.json"
TD = Path(tempfile.mkdtemp())
DST = TD / "datos_hotel.json"
DST.write_bytes(SRC.read_bytes())
os.environ["BM_DEMO_FILE"] = str(DST)

from app.bootstrap import configure_for_flet, get_container, reset_container
from app.core.auth.roles import ROL_DIRECCION
from app.core.auth.session import AuthSession, clear_test_session, set_test_session
from app.core.services import receta_service
from app.core.services.inventory_batch_service import valorizar_cantidad_fifo
from app.core.services.money import as_decimal
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "exports"
OUT.mkdir(parents=True, exist_ok=True)
TODAY = date.today().isoformat()


def _esc(text: object) -> str:
    s = "" if text is None else str(text)
    return s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _fmt(n: float | None, d: int = 4) -> str:
    if n is None:
        return ""
    return f"{n:.{d}f}".rstrip("0").rstrip(".") if d else str(n)


def main() -> None:
    reset_container()
    clear_test_session()
    configure_for_flet()
    set_test_session(
        AuthSession(
            True,
            "usuario",
            "u1",
            "Dir",
            ROL_DIRECCION,
            "s1",
            datetime.now().isoformat(timespec="seconds"),
        )
    )
    data = get_container().app_data_store.get()

    # --- A) Coste unitario por producto (1 unidad de inventario: Kg / Ud / L) ---
    productos_rows: list[dict] = []
    prod_by_id = {}
    for p in sorted(data.productos, key=lambda x: (x.nombre or "").lower()):
        if not getattr(p, "activo", True):
            continue
        uni = p.unidad.value if hasattr(p.unidad, "value") else str(p.unidad)
        # Valorizar 1 unidad nativa vía FIFO
        val = valorizar_cantidad_fifo(data, p.id, 1.0)
        cu = val.coste_unitario_aplicable
        if cu is None and val.coste and not val.incompleto:
            cu = float(val.coste)
        stock = 0.0
        for lote in data.lotes or []:
            if lote.producto_id == p.id and not getattr(lote, "anulado", False):
                stock += float(getattr(lote, "cantidad_restante", 0) or 0)
        row = {
            "producto_id": p.id,
            "nombre": p.nombre,
            "unidad_inventario": uni,
            "codigo": getattr(p, "codigo", None) or "",
            "stock_actual": round(stock, 4),
            "coste_unitario_eur": round(float(cu), 6) if cu is not None else "",
            "coste_completo_1ud": (not val.incompleto) and cu is not None,
            "nota": (
                ""
                if (not val.incompleto and cu is not None)
                else "Sin precio/lote suficiente para 1 ud"
            ),
        }
        productos_rows.append(row)
        prod_by_id[p.id] = row

    # --- B) Porción estándar: cada ingrediente en cada receta ---
    porcion_rows: list[dict] = []
    for rec in sorted(data.recetas or [], key=lambda x: (x.nombre or "").lower()):
        if not getattr(rec, "activo", True):
            continue
        res = receta_service.valorar_receta(rec.id)
        cat = getattr(rec, "categoria", "")
        if hasattr(cat, "value"):
            cat = cat.value
        porc = float(getattr(rec, "porciones_estandar", 1) or 1)
        if not res.ok:
            continue
        for ln in res.lineas or []:
            master = prod_by_id.get(ln.producto_id, {})
            cu_master = master.get("coste_unitario_eur", "")
            # coste de la cantidad de la ración = ya en ln.coste_estimado (factor 1 = porción estándar)
            qty = float(ln.cantidad_nativa or 0)
            coste_porcion = float(ln.coste_estimado or 0)
            # equivalente: si cu conocido, qty * cu
            cu_num = None
            try:
                cu_num = float(cu_master) if cu_master != "" else None
            except (TypeError, ValueError):
                cu_num = None
            coste_calc = round(qty * cu_num, 4) if cu_num is not None else ""
            porcion_rows.append(
                {
                    "receta_id": rec.id,
                    "receta": rec.nombre,
                    "categoria": cat,
                    "porciones_estandar": porc,
                    "producto_id": ln.producto_id,
                    "producto": ln.nombre,
                    "unidad": ln.unidad_nativa,
                    "coste_unitario_producto_eur": cu_master,
                    "cantidad_por_porcion": round(qty, 6),
                    "coste_de_esa_porcion_eur": round(coste_porcion, 4),
                    "check_qty_x_unitario": coste_calc,
                    "coste_incompleto_linea": ln.coste_incompleto,
                    "coste_total_receta_eur": res.coste_total,
                    "coste_por_racion_receta_eur": res.coste_por_racion,
                }
            )

    csv_prod = OUT / f"desglose_coste_unitario_producto_{TODAY}.csv"
    csv_porc = OUT / f"desglose_coste_porcion_estandar_{TODAY}.csv"
    with csv_prod.open("w", newline="", encoding="utf-8-sig") as f:
        w = csv.DictWriter(f, fieldnames=list(productos_rows[0].keys()))
        w.writeheader()
        w.writerows(productos_rows)
    with csv_porc.open("w", newline="", encoding="utf-8-sig") as f:
        w = csv.DictWriter(f, fieldnames=list(porcion_rows[0].keys()))
        w.writeheader()
        w.writerows(porcion_rows)

    # PDF
    pdf_path = OUT / f"desglose_unitario_vs_porcion_{TODAY}.pdf"
    _build_pdf(productos_rows, porcion_rows, pdf_path)
    docx_path = OUT / f"desglose_unitario_vs_porcion_{TODAY}.docx"
    _build_docx(productos_rows, porcion_rows, docx_path)

    # Champiñón sample to console
    print("CSV producto:", csv_prod.resolve())
    print("CSV porcion :", csv_porc.resolve())
    print("PDF         :", pdf_path.resolve())
    print("DOCX        :", docx_path.resolve())
    print("--- Ejemplo CHAMPIÑON ---")
    for r in productos_rows:
        if "champi" in (r["nombre"] or "").lower() and "laminado" in (r["nombre"] or "").lower():
            print("PRODUCTO:", r)
    for r in porcion_rows:
        if "champi" in (r["producto"] or "").lower() and "laminado" in (r["producto"] or "").lower():
            print(
                f"  → {r['receta']}: {r['cantidad_por_porcion']} {r['unidad']} "
                f"× €{r['coste_unitario_producto_eur']}/ud = €{r['coste_de_esa_porcion_eur']} en la ración"
            )


def _build_pdf(productos: list[dict], porciones: list[dict], path: Path) -> None:
    styles = getSampleStyleSheet()
    title = ParagraphStyle("T", parent=styles["Heading1"], fontSize=13, spaceAfter=6)
    h2 = ParagraphStyle("H", parent=styles["Heading2"], fontSize=10, spaceBefore=8, spaceAfter=4)
    body = ParagraphStyle("B", parent=styles["Normal"], fontSize=7.5, leading=9)
    cell = ParagraphStyle("C", parent=styles["Normal"], fontSize=6, leading=7.5)

    doc = SimpleDocTemplate(
        str(path),
        pagesize=landscape(A4),
        leftMargin=8 * mm,
        rightMargin=8 * mm,
        topMargin=8 * mm,
        bottomMargin=8 * mm,
    )
    story: list = []
    story.append(Paragraph("Desglose: coste unitario del producto vs coste por porción", title))
    story.append(
        Paragraph(
            f"Fecha {TODAY}. "
            "1) Coste de 1 Kg / 1 Ud / 1 L del producto en almacén. "
            "2) Cantidad que usa cada receta en 1 porción estándar y su coste. "
            "Ej.: lata champiñón 3 kg = coste de la Ud (envase); en desayuno se carga solo la fracción.",
            body,
        )
    )

    story.append(Paragraph("A) Coste unitario por producto (€ / unidad de inventario)", h2))
    hdr = [
        Paragraph(x, cell)
        for x in (
            "<b>ID</b>",
            "<b>Producto</b>",
            "<b>Ud</b>",
            "<b>Código</b>",
            "<b>Stock</b>",
            "<b>€ / 1 ud</b>",
            "<b>OK</b>",
        )
    ]
    data = [hdr]
    for r in productos:
        data.append(
            [
                Paragraph(_esc(r["producto_id"]), cell),
                Paragraph(_esc(r["nombre"]), cell),
                Paragraph(_esc(r["unidad_inventario"]), cell),
                Paragraph(_esc(r["codigo"]), cell),
                Paragraph(_esc(r["stock_actual"]), cell),
                Paragraph(_esc(r["coste_unitario_eur"]), cell),
                Paragraph("Sí" if r["coste_completo_1ud"] else "No", cell),
            ]
        )
    t = Table(data, colWidths=[16 * mm, 110 * mm, 12 * mm, 26 * mm, 20 * mm, 22 * mm, 12 * mm], repeatRows=1)
    t.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1B3A4B")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("GRID", (0, 0), (-1, -1), 0.2, colors.HexColor("#CCC")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F5F7FA")]),
                ("LEFTPADDING", (0, 0), (-1, -1), 1.5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 1.5),
                ("TOPPADDING", (0, 0), (-1, -1), 1.5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 1.5),
            ]
        )
    )
    story.append(t)

    story.append(Paragraph("B) Coste del producto dentro de cada porción estándar de receta", h2))
    hdr2 = [
        Paragraph(x, cell)
        for x in (
            "<b>Receta</b>",
            "<b>Producto</b>",
            "<b>€/ud prod.</b>",
            "<b>Cant. / porción</b>",
            "<b>Ud</b>",
            "<b>€ en la porción</b>",
            "<b>€/ración receta</b>",
            "<b>Línea OK</b>",
        )
    ]
    data2 = [hdr2]
    for r in porciones:
        data2.append(
            [
                Paragraph(_esc(r["receta"]), cell),
                Paragraph(_esc(r["producto"]), cell),
                Paragraph(_esc(r["coste_unitario_producto_eur"]), cell),
                Paragraph(_esc(r["cantidad_por_porcion"]), cell),
                Paragraph(_esc(r["unidad"]), cell),
                Paragraph(_esc(r["coste_de_esa_porcion_eur"]), cell),
                Paragraph(_esc(r["coste_por_racion_receta_eur"]), cell),
                Paragraph("No" if r["coste_incompleto_linea"] else "Sí", cell),
            ]
        )
    t2 = Table(
        data2,
        colWidths=[42 * mm, 70 * mm, 20 * mm, 22 * mm, 12 * mm, 22 * mm, 22 * mm, 14 * mm],
        repeatRows=1,
    )
    t2.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#0F6E56")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("GRID", (0, 0), (-1, -1), 0.2, colors.HexColor("#CCC")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F3FAF7")]),
                ("LEFTPADDING", (0, 0), (-1, -1), 1.5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 1.5),
                ("TOPPADDING", (0, 0), (-1, -1), 1.5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 1.5),
            ]
        )
    )
    story.append(t2)

    # Champiñón callout
    story.append(Spacer(1, 6))
    story.append(Paragraph("Ejemplo champiñón laminado 3 kg", h2))
    for r in productos:
        if "champi" in (r["nombre"] or "").lower() and "laminado" in (r["nombre"] or "").lower():
            story.append(
                Paragraph(
                    f"Producto: <b>{_esc(r['nombre'])}</b> — coste de "
                    f"<b>1 {_esc(r['unidad_inventario'])}</b> (envase/lata) = "
                    f"<b>{_esc(r['coste_unitario_eur'])} €</b>.",
                    body,
                )
            )
    for r in porciones:
        if "champi" in (r["producto"] or "").lower() and "laminado" in (r["producto"] or "").lower():
            story.append(
                Paragraph(
                    f"· Receta <b>{_esc(r['receta'])}</b>: usa {_esc(r['cantidad_por_porcion'])} {_esc(r['unidad'])} "
                    f"→ coste en la porción <b>{_esc(r['coste_de_esa_porcion_eur'])} €</b> "
                    f"(ración completa receta {_esc(r['coste_por_racion_receta_eur'])} €).",
                    body,
                )
            )

    doc.build(story)


def _build_docx(productos: list[dict], porciones: list[dict], path: Path) -> None:
    document = Document()
    section = document.sections[0]
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.left_margin = Cm(1.0)
    section.right_margin = Cm(1.0)
    section.top_margin = Cm(1.0)
    section.bottom_margin = Cm(1.0)

    document.add_heading("Desglose: coste unitario vs coste por porción", level=1)
    p = document.add_paragraph(
        f"Fecha {TODAY}. Tabla A = € por 1 Kg/Ud/L del producto. "
        "Tabla B = cantidad y € de ese producto en 1 porción estándar de cada receta."
    )
    for run in p.runs:
        run.font.size = Pt(9)

    document.add_heading("A) Coste unitario por producto", level=2)
    t = document.add_table(rows=1, cols=7)
    t.style = "Table Grid"
    for i, h in enumerate(["ID", "Producto", "Ud", "Código", "Stock", "€ / 1 ud", "OK"]):
        t.rows[0].cells[i].text = h
    for r in productos:
        row = t.add_row().cells
        row[0].text = r["producto_id"]
        row[1].text = r["nombre"] or ""
        row[2].text = r["unidad_inventario"] or ""
        row[3].text = r["codigo"] or ""
        row[4].text = str(r["stock_actual"])
        row[5].text = str(r["coste_unitario_eur"])
        row[6].text = "Sí" if r["coste_completo_1ud"] else "No"

    document.add_heading("B) Coste en la porción estándar de cada receta", level=2)
    t2 = document.add_table(rows=1, cols=8)
    t2.style = "Table Grid"
    for i, h in enumerate(
        [
            "Receta",
            "Producto",
            "€/ud prod.",
            "Cant./porción",
            "Ud",
            "€ en porción",
            "€/ración receta",
            "Línea OK",
        ]
    ):
        t2.rows[0].cells[i].text = h
    for r in porciones:
        row = t2.add_row().cells
        row[0].text = r["receta"] or ""
        row[1].text = r["producto"] or ""
        row[2].text = str(r["coste_unitario_producto_eur"])
        row[3].text = str(r["cantidad_por_porcion"])
        row[4].text = r["unidad"] or ""
        row[5].text = str(r["coste_de_esa_porcion_eur"])
        row[6].text = str(r["coste_por_racion_receta_eur"] or "")
        row[7].text = "No" if r["coste_incompleto_linea"] else "Sí"

    document.add_heading("Ejemplo champiñón", level=2)
    for r in productos:
        if "champi" in (r["nombre"] or "").lower() and "laminado" in (r["nombre"] or "").lower():
            document.add_paragraph(
                f"{r['nombre']}: 1 {r['unidad_inventario']} (lata/envase) = {r['coste_unitario_eur']} €"
            )
    for r in porciones:
        if "champi" in (r["producto"] or "").lower() and "laminado" in (r["producto"] or "").lower():
            document.add_paragraph(
                f"{r['receta']}: {r['cantidad_por_porcion']} {r['unidad']} → "
                f"{r['coste_de_esa_porcion_eur']} € en la porción"
            )

    document.save(str(path))


if __name__ == "__main__":
    main()
