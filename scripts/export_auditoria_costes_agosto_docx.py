"""Genera Word imprimible de auditoría de costes del import agosto 2026.

Uso:
  .\\.venv\\Scripts\\python.exe scripts\\export_auditoria_costes_agosto_docx.py
  .\\.venv\\Scripts\\python.exe scripts\\export_auditoria_costes_agosto_docx.py --suffix _post_correccion
"""
from __future__ import annotations

import argparse
import os
import sys
from collections import defaultdict
from datetime import date, datetime
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from app.bootstrap import configure_for_flet, get_container, reset_container

HOTEL = Path(os.environ["LOCALAPPDATA"]) / "BM-V2-local" / "data" / "datos_hotel.json"
OUT_DIR = ROOT / "docs" / "añadidos manual"

CLAVE_PREFIX = "import-ago26"
FECHA_INI = date(2026, 8, 1)
FECHA_FIN = date(2026, 8, 19)

# Packs conocidos (kg netos por Ud) — mismos que seed tostada / corrección
PACK_KG: dict[str, float] = {
    "p117": 3.0,   # champiñón laminado 3KG
    "p185": 0.2,   # espinacas bolsa ~200g
    "p286": 3.0,   # alubia tomate lata 3KG
    "p122": 0.85,  # melocotón en su jugo
    "p405": 2.5,
}

# Extras típicos del parser (gr) usados para estimar qty esperada
EXTRAS_GR_TIPICOS: dict[str, float] = {
    "p117": 20.0,   # champi por ración
    "p185": 15.0,   # espinaca
    "p286": 20.0,   # judías/alubia
    "p122": 375.0,  # melocotón almíbar buffet
}

# Productos Ud de buffet que sí van en piezas (no marcar REVISAR por qty alta)
BUFFET_PIEZAS_OK = {
    "p05", "p357", "p276", "p252", "p251", "p294", "p250", "p249", "b01", "p66", "p48",
}


def _boot():
    reset_container()
    configure_for_flet(data_path=str(HOTEL))
    return get_container().app_data_store.get()


def _is_import(clave: str | None) -> bool:
    return bool(clave and clave.startswith(CLAVE_PREFIX))


def _set_cell_shading(cell, hex_color: str) -> None:
    from docx.oxml import OxmlElement

    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")


def _add_table(doc: Document, headers: list[str], rows: list[list[str]], *, highlight_col: int | None = None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
        _set_cell_shading(hdr[i], "D9E2F3")
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            for p in cells[ci].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(8)
            if highlight_col is not None and ci == highlight_col and str(val).upper() == "REVISAR":
                _set_cell_shading(cells[ci], "F8CBAD")
    doc.add_paragraph()


def collect(data) -> dict:
    by_prod = {p.id: p for p in data.productos}

    # Aggregates
    cost_by_prod: dict[str, float] = defaultdict(float)
    qty_by_prod: dict[str, float] = defaultdict(float)
    cost_by_day_svc: dict[tuple[date, str], float] = defaultdict(float)
    lines_by_day_svc: dict[tuple[date, str], int] = defaultdict(int)
    totals = {
        "desayuno_huesped": 0.0,
        "buffet": 0.0,
        "comida": 0.0,
        "bebidas": 0.0,
        "merma_hielo": 0.0,
    }

    for d in data.desayunos:
        if getattr(d, "anulado", False):
            continue
        clave = getattr(d, "clave_idempotencia", None) or ""
        if not _is_import(clave):
            continue
        if not (FECHA_INI <= d.fecha <= FECHA_FIN):
            continue
        tipo = "buffet" if "buffet" in clave else "desayuno_huesped"
        totals[tipo] += float(d.coste_total or 0)
        cost_by_day_svc[(d.fecha, tipo)] += float(d.coste_total or 0)
        nlin = len(getattr(d, "lineas", None) or [])
        lines_by_day_svc[(d.fecha, tipo)] += nlin
        for lin in d.lineas or []:
            cost_by_prod[lin.producto_id] += float(lin.coste or 0)
            qty_by_prod[lin.producto_id] += float(lin.cantidad or 0)

    for r in data.registros_servicio:
        if getattr(r, "anulado", False):
            continue
        clave = getattr(r, "clave_idempotencia", None) or ""
        if not _is_import(clave):
            continue
        if not (FECHA_INI <= r.fecha <= FECHA_FIN):
            continue
        tipo = r.tipo_servicio  # comida | bebidas
        if tipo not in totals:
            totals[tipo] = 0.0
        totals[tipo] += float(r.coste_total or 0)
        cost_by_day_svc[(r.fecha, tipo)] += float(r.coste_total or 0)
        nlin = len(getattr(r, "lineas", None) or [])
        lines_by_day_svc[(r.fecha, tipo)] += nlin
        for lin in r.lineas or []:
            cost_by_prod[lin.producto_id] += float(lin.coste or 0)
            qty_by_prod[lin.producto_id] += float(lin.cantidad or 0)

    for m in getattr(data, "mermas", None) or []:
        if getattr(m, "anulado", False):
            continue
        if not (FECHA_INI <= m.fecha <= FECHA_FIN):
            continue
        for lin in m.lineas or []:
            com = getattr(lin, "comentario", None) or ""
            if CLAVE_PREFIX in com and "merma-hielo" in com:
                totals["merma_hielo"] += float(lin.coste or 0)
                cost_by_prod[lin.producto_id] += float(lin.coste or 0)
                qty_by_prod[lin.producto_id] += float(lin.cantidad or 0)

    total_all = sum(totals.values())

    ranking = []
    for pid, coste in sorted(cost_by_prod.items(), key=lambda x: -x[1]):
        p = by_prod.get(pid)
        qty = qty_by_prod[pid]
        unidad = p.unidad.value if p else "?"
        nombre = p.nombre if p else pid
        eur_u = (coste / qty) if qty else 0.0
        revisar = False
        if p and p.unidad.value == "Ud" and pid not in BUFFET_PIEZAS_OK:
            # qty total mes absurda o inflación pack
            if qty > 30:
                revisar = True
            pack = PACK_KG.get(pid)
            tip = EXTRAS_GR_TIPICOS.get(pid)
            if pack and tip and tip > 0:
                # estimación muy grosera: si qty >> lo que cabría en ~200 raciones*tip/pack
                esperado_max = 400 * (tip / 1000.0) / pack  # tope holgado
                if qty > max(esperado_max * 10, 5):
                    revisar = True
        if total_all > 0 and coste / total_all > 0.05:
            revisar = True
        ranking.append(
            {
                "pid": pid,
                "nombre": nombre,
                "unidad": unidad,
                "qty": qty,
                "coste": coste,
                "eur_u": eur_u,
                "revisar": revisar,
                "pct": (100 * coste / total_all) if total_all else 0,
            }
        )

    anomalias = []
    for pid, pack in PACK_KG.items():
        tip = EXTRAS_GR_TIPICOS.get(pid, 20.0)
        esperado_por_extra = (tip / 1000.0) / pack if pack else 0
        real = qty_by_prod.get(pid, 0.0)
        # nº aproximado de "extras" si cada uno fuera tip gr mal convertidos a tip Ud
        if tip > 0 and real > 0:
            # si bug: cada extra aporta `tip` Ud; si ok: aporta tip/1000/pack Ud
            extras_si_bug = real / tip
            extras_si_ok = real / esperado_por_extra if esperado_por_extra else 0
            factor = (tip / esperado_por_extra) if esperado_por_extra else 0  # tip Ud / (tip/pack_g)
            # factor teórico bug: tip / (tip/1000/pack) = pack*1000
            factor_teorico = pack * 1000.0 / 1.0  # gr→Ud sin pack: 20 gr → 20 Ud vs 20/3000
            factor_real = real / max(esperado_por_extra * max(extras_si_bug, 1), 1e-9)
            # más útil: qty_esperada si hubiéramos convertido bien el mismo nº de extras mal contados
            n_extras_estim = real / tip if tip else 0  # asume bug clásico
            qty_esperada = n_extras_estim * esperado_por_extra
            inflacion = (real / qty_esperada) if qty_esperada > 1e-12 else 0
        else:
            n_extras_estim = 0
            qty_esperada = 0
            inflacion = 0
            factor_teorico = pack * 1000.0
        p = by_prod.get(pid)
        anomalias.append(
            {
                "pid": pid,
                "nombre": p.nombre if p else pid,
                "pack_kg": pack,
                "extra_gr": tip,
                "qty_real": real,
                "qty_esperada": qty_esperada,
                "inflacion": inflacion,
                "coste": cost_by_prod.get(pid, 0.0),
                "factor_teorico": factor_teorico,
            }
        )

    repo_lots = [
        l
        for l in data.lotes
        if getattr(l, "marca_proveedor", None) in ("IMPORT-AGO26", "BUFFET-AGO26")
    ]

    return {
        "totals": totals,
        "total_all": total_all,
        "ranking": ranking,
        "anomalias": anomalias,
        "cost_by_day_svc": cost_by_day_svc,
        "lines_by_day_svc": lines_by_day_svc,
        "by_prod": by_prod,
        "repo_lots": repo_lots,
    }


def build_docx(data_view: dict, out_path: Path, *, titulo_extra: str = "") -> None:
    doc = Document()
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)

    # Portada
    t = doc.add_heading("Auditoría de costes — Import agosto 2026", level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if titulo_extra:
        p = doc.add_paragraph(titulo_extra)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Periodo: {FECHA_INI.strftime('%d/%m/%Y')} – {FECHA_FIN.strftime('%d/%m/%Y')}")
    doc.add_paragraph(f"Generado: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    doc.add_paragraph(
        "Claves incluidas: registros con clave_idempotencia que empieza por "
        f"«{CLAVE_PREFIX}» (desayuno huésped, buffet, comida, bebidas) y mermas de hielo asociadas."
    )

    aviso = doc.add_paragraph()
    run = aviso.add_run(
        "AVISO: hay errores de conversión gr→Ud en productos envasados (bolsa/lata). "
        "Los importes de champiñón, alubia y espinacas están inflados y NO son realistas. "
        "Use este documento para revisar a mano antes/después de la corrección."
    )
    run.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    doc.add_heading("1. Totales por tipo", level=1)
    tot = data_view["totals"]
    _add_table(
        doc,
        ["Tipo", "Coste (€)"],
        [
            ["Desayuno huésped", f"{tot['desayuno_huesped']:.2f}"],
            ["Buffet diario", f"{tot['buffet']:.2f}"],
            ["Comida TPV", f"{tot.get('comida', 0):.2f}"],
            ["Bebidas TPV", f"{tot.get('bebidas', 0):.2f}"],
            ["Merma hielo", f"{tot['merma_hielo']:.2f}"],
            ["TOTAL", f"{data_view['total_all']:.2f}"],
        ],
    )

    repo = data_view["repo_lots"]
    doc.add_paragraph(
        f"Lotes de reposición automática IMPORT-AGO26 / BUFFET-AGO26: {len(repo)} "
        f"(precio total entradas ≈ {sum(float(l.precio_total or 0) for l in repo):.2f} €)."
    )

    # Ranking
    doc.add_heading("2. Ranking productos (top 30 por coste)", level=1)
    doc.add_paragraph(
        "Marca REVISAR si: Ud (no bollería buffet) con qty alta, inflación pack, o >5 % del total."
    )
    rows = []
    for i, r in enumerate(data_view["ranking"][:30], 1):
        rows.append(
            [
                str(i),
                r["pid"],
                (r["nombre"][:42] + "…") if len(r["nombre"]) > 43 else r["nombre"],
                r["unidad"],
                f"{r['qty']:.3f}",
                f"{r['coste']:.2f}",
                f"{r['eur_u']:.2f}",
                f"{r['pct']:.1f}%",
                "REVISAR" if r["revisar"] else "ok",
            ]
        )
    _add_table(
        doc,
        ["#", "Id", "Producto", "Ud", "Qty", "Coste €", "€/u", "%", "Flag"],
        rows,
        highlight_col=8,
    )

    # Anomalías
    doc.add_heading("3. Anomalías Ud-pack (gr mal convertidos a Ud)", level=1)
    doc.add_paragraph(
        "Si el bug aplica, 20 g de un pack de 3 kg se registraron como 20 Ud (bolsas) "
        "en lugar de 20/3000 ≈ 0,0067 Ud. Factor teórico ≈ pack_kg × 1000."
    )
    arows = []
    for a in data_view["anomalias"]:
        arows.append(
            [
                a["pid"],
                (a["nombre"][:36] + "…") if len(a["nombre"]) > 37 else a["nombre"],
                f"{a['pack_kg']:g} kg",
                f"{a['extra_gr']:g} g",
                f"{a['qty_real']:.3f}",
                f"{a['qty_esperada']:.4f}",
                f"{a['inflacion']:.0f}×" if a["inflacion"] else "—",
                f"{a['coste']:.2f}",
                f"~{a['factor_teorico']:.0f}×",
            ]
        )
    _add_table(
        doc,
        ["Id", "Producto", "Pack", "Extra tip.", "Qty real", "Qty esp.", "Inflación", "Coste €", "Factor teor."],
        arows,
    )

    # Desglose diario
    doc.add_heading("4. Desglose por día × servicio", level=1)
    doc.add_paragraph("Puede marcar a boli la columna OK / ?.")
    drows = []
    day = FECHA_INI
    while day <= FECHA_FIN:
        for svc in ("desayuno_huesped", "buffet", "comida", "bebidas"):
            c = data_view["cost_by_day_svc"].get((day, svc), 0.0)
            n = data_view["lines_by_day_svc"].get((day, svc), 0)
            if c == 0 and n == 0:
                continue
            label = {
                "desayuno_huesped": "Des. huésped",
                "buffet": "Buffet",
                "comida": "Comida",
                "bebidas": "Bebidas",
            }[svc]
            drows.append(
                [
                    day.strftime("%d/%m"),
                    label,
                    f"{c:.2f}",
                    str(n),
                    "",  # OK/?
                    "",  # notas
                ]
            )
        day = date.fromordinal(day.toordinal() + 1)
    _add_table(
        doc,
        ["Fecha", "Servicio", "Coste €", "Líneas", "OK/?", "Notas a mano"],
        drows,
    )

    # Notas / checklist
    doc.add_heading("5. Checklist de revisión manual", level=1)
    doc.add_paragraph("Anote aquí lo que haya que modificar y devuelva el documento (foto o lista):")
    checks = [
        "[ ] Champiñón p117 — corregir conversión gr→Ud (pack 3 kg)",
        "[ ] Alubia p286 — corregir conversión gr→Ud (pack 3 kg)",
        "[ ] Espinacas p185 — corregir conversión (pack ~0,2 kg)",
        "[ ] Otros productos Ud con REVISAR: ___________________________",
        "[ ] Buffet panes/bollería — cantidades OK / ajustar: ___________",
        "[ ] TPV comida/bebidas — parece OK / revisar días: _____________",
        "[ ] Anular import-ago26 y reimportar con PACK_KG",
        "[ ] Limpiar lotes fantasma IMPORT-AGO26 / BUFFET-AGO26",
        "[ ] Yogur / pan masa madre (pendientes de SKU)",
        "[ ] Cóctel del día (lista pendiente)",
    ]
    for c in checks:
        doc.add_paragraph(c)
    doc.add_paragraph()
    doc.add_paragraph("Observaciones libres:")
    for _ in range(8):
        doc.add_paragraph("_" * 85)

    doc.add_paragraph()
    pie = doc.add_paragraph(
        "Documento generado automáticamente desde datos_hotel.json — solo lectura; "
        "no modifica stock ni registros."
    )
    for r in pie.runs:
        r.italic = True
        r.font.size = Pt(8)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("--suffix", default="", help="Sufijo de fichero, p.ej. _post_correccion")
    ap.add_argument("--titulo", default="", help="Subtítulo en portada")
    args = ap.parse_args()

    data = _boot()
    view = collect(data)
    name = f"Auditoria_costes_agosto_2026{args.suffix}.docx"
    out = OUT_DIR / name
    titulo = args.titulo or ("Post-corrección" if args.suffix else "Pre-corrección — revisar a mano")
    build_docx(view, out, titulo_extra=titulo)
    print("Escrito:", out)
    print("Total import:", round(view["total_all"], 2), "EUR")
    print("Top3 REVISAR:", [(r["pid"], round(r["coste"], 2)) for r in view["ranking"] if r["revisar"]][:5])
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
