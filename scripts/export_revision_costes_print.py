"""Genera PDF y Word imprimibles desde los CSV de productos/recetas."""

from __future__ import annotations

import csv
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import (
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

ROOT = Path(__file__).resolve().parents[1]
EXPORTS = ROOT / "docs" / "exports"
PROD_CSV = EXPORTS / "productos_stock_coste.csv"
REC_CSV = EXPORTS / "recetas_coste_por_racion.csv"
OUT_PDF = EXPORTS / f"revision_costes_{date.today().isoformat()}.pdf"
OUT_DOCX = EXPORTS / f"revision_costes_{date.today().isoformat()}.docx"


def _read_csv(path: Path) -> list[dict[str, str]]:
    with path.open(encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def _esc(text: object) -> str:
    s = "" if text is None else str(text)
    return (
        s.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )


def build_pdf(productos: list[dict], recetas: list[dict], path: Path) -> None:
    styles = getSampleStyleSheet()
    title = ParagraphStyle(
        "TitleES",
        parent=styles["Heading1"],
        fontSize=14,
        spaceAfter=8,
    )
    h2 = ParagraphStyle(
        "H2ES",
        parent=styles["Heading2"],
        fontSize=11,
        spaceBefore=10,
        spaceAfter=6,
    )
    body = ParagraphStyle(
        "BodyES",
        parent=styles["Normal"],
        fontSize=8,
        leading=10,
    )
    cell = ParagraphStyle(
        "CellES",
        parent=styles["Normal"],
        fontSize=6.5,
        leading=8,
    )

    doc = SimpleDocTemplate(
        str(path),
        pagesize=landscape(A4),
        leftMargin=10 * mm,
        rightMargin=10 * mm,
        topMargin=10 * mm,
        bottomMargin=10 * mm,
        title="Revisión productos y costes por ración",
    )
    story: list = []
    story.append(Paragraph("Revisión de productos, stock y costes", title))
    story.append(
        Paragraph(
            f"Fecha: {date.today().isoformat()} · Fuente: datos_hotel local · "
            f"{len(productos)} productos · {len(recetas)} recetas · "
            "Costes estimados FIFO/lote (revisar incompletos).",
            body,
        )
    )
    story.append(Spacer(1, 6))

    # --- Productos ---
    story.append(Paragraph("1. Productos — stock y coste unitario", h2))
    prod_header = [
        Paragraph("<b>ID</b>", cell),
        Paragraph("<b>Nombre</b>", cell),
        Paragraph("<b>Ud</b>", cell),
        Paragraph("<b>Código</b>", cell),
        Paragraph("<b>Stock</b>", cell),
        Paragraph("<b>€/ud est.</b>", cell),
        Paragraph("<b>Valor stock</b>", cell),
    ]
    prod_data = [prod_header]
    for r in productos:
        prod_data.append(
            [
                Paragraph(_esc(r.get("producto_id")), cell),
                Paragraph(_esc(r.get("nombre")), cell),
                Paragraph(_esc(r.get("unidad")), cell),
                Paragraph(_esc(r.get("codigo")), cell),
                Paragraph(_esc(r.get("stock_restante_lotes")), cell),
                Paragraph(_esc(r.get("coste_unitario_estimado")), cell),
                Paragraph(_esc(r.get("valor_stock_estimado")), cell),
            ]
        )
    t1 = Table(
        prod_data,
        colWidths=[18 * mm, 95 * mm, 14 * mm, 28 * mm, 22 * mm, 22 * mm, 24 * mm],
        repeatRows=1,
    )
    t1.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1B3A4B")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#CCCCCC")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F5F7FA")]),
                ("LEFTPADDING", (0, 0), (-1, -1), 2),
                ("RIGHTPADDING", (0, 0), (-1, -1), 2),
                ("TOPPADDING", (0, 0), (-1, -1), 2),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
            ]
        )
    )
    story.append(t1)

    # --- Recetas ---
    story.append(Paragraph("2. Recetas — coste por ración", h2))
    rec_header = [
        Paragraph("<b>ID</b>", cell),
        Paragraph("<b>Receta</b>", cell),
        Paragraph("<b>Cat.</b>", cell),
        Paragraph("<b>Porc.</b>", cell),
        Paragraph("<b>€ total</b>", cell),
        Paragraph("<b>€/ración</b>", cell),
        Paragraph("<b>Completo</b>", cell),
        Paragraph("<b>Desglose ingredientes</b>", cell),
    ]
    rec_data = [rec_header]
    for r in recetas:
        completo = r.get("coste_completo", "")
        flag = "Sí" if str(completo).lower() == "true" else "No"
        rec_data.append(
            [
                Paragraph(_esc(r.get("receta_id")), cell),
                Paragraph(_esc(r.get("nombre")), cell),
                Paragraph(_esc(r.get("categoria")), cell),
                Paragraph(_esc(r.get("porciones_estandar")), cell),
                Paragraph(_esc(r.get("coste_total")), cell),
                Paragraph(_esc(r.get("coste_por_racion")), cell),
                Paragraph(flag, cell),
                Paragraph(_esc(r.get("desglose_ingredientes")), cell),
            ]
        )
    t2 = Table(
        rec_data,
        colWidths=[14 * mm, 40 * mm, 18 * mm, 14 * mm, 16 * mm, 16 * mm, 16 * mm, 110 * mm],
        repeatRows=1,
    )
    t2.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#0F6E56")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#CCCCCC")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F3FAF7")]),
                ("LEFTPADDING", (0, 0), (-1, -1), 2),
                ("RIGHTPADDING", (0, 0), (-1, -1), 2),
                ("TOPPADDING", (0, 0), (-1, -1), 2),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
            ]
        )
    )
    story.append(t2)

    # Huevos summary page note
    story.append(Spacer(1, 8))
    story.append(Paragraph("Nota huevos (revisión)", h2))
    huevos = [r for r in productos if "huevo" in (r.get("nombre") or "").lower()]
    for h in huevos:
        story.append(
            Paragraph(
                f"· {_esc(h.get('nombre'))} ({_esc(h.get('producto_id'))}): "
                f"stock {_esc(h.get('stock_restante_lotes'))} {_esc(h.get('unidad'))}, "
                f"€/ud {_esc(h.get('coste_unitario_estimado'))}",
                body,
            )
        )
    story.append(
        Paragraph(
            "No aparece stock de ~357.000 ud de huevos en lotes. "
            "HUEVO CÁSCARA puede mostrar descuadre lote vs ledger; revisar movimientos.",
            body,
        )
    )

    doc.build(story)


def build_docx(productos: list[dict], recetas: list[dict], path: Path) -> None:
    document = Document()
    section = document.sections[0]
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)

    title = document.add_heading("Revisión de productos, stock y costes", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p = document.add_paragraph(
        f"Fecha: {date.today().isoformat()} · {len(productos)} productos · "
        f"{len(recetas)} recetas · Costes estimados (FIFO/lote)."
    )
    for run in p.runs:
        run.font.size = Pt(9)

    document.add_heading("1. Productos — stock y coste unitario", level=2)
    table = document.add_table(rows=1, cols=7)
    table.style = "Table Grid"
    headers = ["ID", "Nombre", "Ud", "Código", "Stock", "€/ud est.", "Valor stock"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for r in productos:
        row = table.add_row().cells
        row[0].text = r.get("producto_id") or ""
        row[1].text = r.get("nombre") or ""
        row[2].text = r.get("unidad") or ""
        row[3].text = r.get("codigo") or ""
        row[4].text = r.get("stock_restante_lotes") or ""
        row[5].text = r.get("coste_unitario_estimado") or ""
        row[6].text = r.get("valor_stock_estimado") or ""

    document.add_heading("2. Recetas — coste por ración", level=2)
    table2 = document.add_table(rows=1, cols=8)
    table2.style = "Table Grid"
    headers2 = [
        "ID",
        "Receta",
        "Cat.",
        "Porc.",
        "€ total",
        "€/ración",
        "Completo",
        "Desglose",
    ]
    for i, h in enumerate(headers2):
        table2.rows[0].cells[i].text = h
    for r in recetas:
        completo = r.get("coste_completo", "")
        flag = "Sí" if str(completo).lower() == "true" else "No"
        row = table2.add_row().cells
        row[0].text = r.get("receta_id") or ""
        row[1].text = r.get("nombre") or ""
        row[2].text = r.get("categoria") or ""
        row[3].text = str(r.get("porciones_estandar") or "")
        row[4].text = str(r.get("coste_total") or "")
        row[5].text = str(r.get("coste_por_racion") or "")
        row[6].text = flag
        row[7].text = r.get("desglose_ingredientes") or ""

    document.add_heading("Nota huevos", level=2)
    for h in (r for r in productos if "huevo" in (r.get("nombre") or "").lower()):
        document.add_paragraph(
            f"{h.get('nombre')} ({h.get('producto_id')}): "
            f"stock {h.get('stock_restante_lotes')} {h.get('unidad')}, "
            f"€/ud {h.get('coste_unitario_estimado')}"
        )
    document.add_paragraph(
        "No hay ~357.000 ud de huevos en lotes. Revisar descuadre ledger vs restante en HUEVO CÁSCARA."
    )

    document.save(str(path))


def main() -> None:
    if not PROD_CSV.exists() or not REC_CSV.exists():
        raise SystemExit(
            f"Faltan CSV. Ejecuta antes scripts/export_productos_recetas_coste.py\n"
            f"  {PROD_CSV}\n  {REC_CSV}"
        )
    productos = _read_csv(PROD_CSV)
    recetas = _read_csv(REC_CSV)
    build_pdf(productos, recetas, OUT_PDF)
    build_docx(productos, recetas, OUT_DOCX)
    print("PDF ", OUT_PDF.resolve())
    print("DOCX", OUT_DOCX.resolve())


if __name__ == "__main__":
    main()
